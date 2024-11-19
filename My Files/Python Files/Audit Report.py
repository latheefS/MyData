'''
@module_name :       Audit Report
@author :            Shaik Latheef
@created_date :      09 May 2024
'''

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import docx
import traceback
import warnings
warnings.simplefilter('ignore')

def mark_others(row):
    if row['Result'] != 'Passed' and row['Result'] != 'Failed' and row['Result'] != 'N/A':
        row['Result'] = 'Others'
    return row


def add_test_lists(doc,df):
    try:
        # Add a Title to the document
        doc.add_heading('Test Scenarios List', level=1)
        list_df = df[["Test Id", "Test Name", "Test Area"]]
        df_rows = tuple(list_df.itertuples(index=False, name=None))
        # Creating a table object
        table = doc.add_table(rows=1, cols=len(list_df.columns))
        table.style = 'TableGrid'

        # Adding heading in the 1st row of the table
        row = table.rows[0].cells
        cols = list(list_df.columns)
        for i in range(len(cols)):
            row[i].text = cols[i]

        # Adding data from the list to the table
        for id, name, area in df_rows:
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(id)
            row[1].text = name
            row[2].text = area
        return doc
    except:
        print(f"Exception Occured:\n{traceback.format_exc()}")

def add_failed_tests(doc,df):
    try:
        # Generate list of tests failed
        test_failed_df = df[df["Result"] == 'Failed']
        test_failed_df = test_failed_df[["Test Id", "Test Name", "Failed Step", "Reasons", "Error Message"]]

        df_rows = tuple(test_failed_df.itertuples(index=False, name=None))
        # Creating a table object
        doc.add_heading('List of tests failed', level=1)
        table = doc.add_table(rows=1, cols=len(test_failed_df.columns))
        table.style = 'TableGrid'

        # Adding heading in the 1st row of the table
        row = table.rows[0].cells

        cols = list(test_failed_df.columns)
        for i in range(len(cols)):
            row[i].text = cols[i]

        # Adding data from the list to the table

        for id, name, failed_step, reasons, error_message in df_rows:
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(id)
            row[1].text = str(name)
            row[2].text = str(failed_step)
            row[3].text = str(reasons)
            row[4].text = str(error_message)
        return doc
    except:
        print(f"Exception Occured:\n{traceback.format_exc()}")


try:
    '''
        Reads the input csv file as a dataframe and saves the data to word document as
        Table with list of Test Scenarios,
        Pie chart depicting Passed and Failed tests,
        Bar chart depicting Passed and Failed tests,
        Table with list of Failed tests.
    '''
    # Prompt the user for the CSV file path
    csv_file_path = input("Enter the CSV file path: ")
    # csv_file_path = 'C:/Users/LatheefS/Desktop/AT Python Files/x_ve1_maximise_test_results.csv'

    # Read the CSV file into a pandas DataFrame
    df = pd.read_csv(csv_file_path)
    df = df.fillna('')
    df.rename(columns={'number': 'Number',
                       'u_company': 'Company',
                       'error_message': 'Error Message',
                       'failed_step': 'Failed Step',
                       'process_id': 'Process Id',
                       'reasons': 'Reasons',
                       'result': 'Result',
                       'test_area': 'Test Area',
                       'test_duration': 'Test Duration',
                       'test_id': 'Test Id',
                       'test_name': 'Test Name'}, inplace=True)
    # Add tests list table to document
    doc = Document()
    doc = add_test_lists(doc,df)
    pie_df = df.copy()
    pie_df = pie_df.apply(mark_others, axis=1)

    # Generate Pie Chart
    # Assuming your 'Result' column contains distinct values (e.g., 'Passed', 'Failed', etc.)
    result_counts = pie_df['Result'].value_counts()
    # Define custom colors for each value
    colors = {'Passed': '#7bdcb5','Failed': '#f78da7','Others': '#fcb900'}         # Add more colors for other values if needed
    # Create the donut pie chart
    plt.figure(figsize=(8, 6))
    # Calculate the total count and plot the center text (total count)
    total_count = result_counts.sum()
    plt.text(0, 0, f"{total_count}", fontsize=12, ha='center', va='center')
    # Create the pie chart
    wedges, _, _ = plt.pie(result_counts, labels=result_counts.index,
                           colors=[colors.get(val, '#abb8c3') for val in result_counts.index],
                           autopct='', pctdistance=0.85, wedgeprops={'edgecolor': 'w'})
    # Add individual counts on each wedge
    for wedge, count in zip(wedges, result_counts):
        angle = (wedge.theta2 - wedge.theta1) / 2.0 + wedge.theta1
        x = 0.85 * np.cos(np.deg2rad(angle))
        y = 0.85 * np.sin(np.deg2rad(angle))
        plt.text(x, y, str(count), ha='center', va='center')

    # Draw a white circle at the center to create the donut effect
    centre_circle = plt.Circle((0, 0), 0.50, fc='white')
    fig = plt.gcf()
    fig.gca().add_artist(centre_circle)

    # Add title and legend
    plt.title(r"$\bf{Release\ Overview}$", color='brown')
    plt.legend(result_counts.index, loc='upper right', prop={'weight': 'bold'})
    plt.axis('equal')  # Equal aspect ratio ensures a circular pie chart
    # plt.show() # To display during execution
    # Save the Pie chart to an image file
    piechart_image_path = 'pie_chart.png'
    plt.savefig(piechart_image_path)

    # Print the chart to a Word document
    doc.add_heading('3.Test Overview', level=1)
    doc.add_picture(piechart_image_path, width=docx.shared.Inches(6))
    print("Pie Chart saved as 'pie_chart.png' and printed to 'AT Completion Report.docx'.")
    plt.close()

    # Generate Bar Chart
    color_map = {'Passed':'#7bdcb5','Failed':'#f78da7'}
    bar_df = df.copy()
    bar_df = bar_df[(bar_df['Result']=='Passed')|(bar_df['Result']=='Failed')]
    bargraph = sns.countplot(x='Test Area', hue='Result', data=bar_df, palette=color_map)
    sns.set(rc={'figure.figsize': (10, 10)})
    for i in bargraph.containers:
        bargraph.bar_label(i, )

    # Customize plot properties
    plt.title("Results By Test Area", color='brown', fontsize=14, weight='bold')
    plt.xlabel("Test Area")
    plt.ylabel("No.of Tests")
    plt.legend(loc='upper right', prop={'weight': 'bold'})
    # plt.show() # To display during execution

    # Save the Bar Graph to an image file
    bargraph_image_path = 'bar_chart.png'
    plt.savefig(bargraph_image_path)
    # Print the chart to a Word document
    doc.add_picture(bargraph_image_path, width=docx.shared.Inches(6))
    print("Bar Chart saved as 'bar_chart.png' and printed to 'AT Completion Report.docx'.")


    # Add failed tests list table to document and save the document
    doc = add_failed_tests(doc, df)
    doc.save('AT Completion Report.docx')


except FileNotFoundError:
    print(f"File '{csv_file_path}' not found. Please check the path and try again.")
    exit()

except:
    print(f"Exception occurred:{traceback.format_exc()}")