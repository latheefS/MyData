'''
@module_name :       Audit Report
@author :            Shaik Latheef
@created_date :      09 May 2024
'''

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import docx
from docx import Document
import traceback
import warnings
import Static
from docx.shared import Pt

warnings.simplefilter('ignore')
import PyAi

def mark_others(row):
    if row['Result'] != 'Passed' and row['Result'] != 'Failed' and row['Result'] != 'N/A':
        row['Result'] = 'Others'
    return row


def add_test_lists(doc,df):
    try:
        # Add a Title to the document
        #doc.add_heading('Test Scenarios List', level=1)
        list_df = df[["Test Id", "Test Name", "Test Area"]]
        df_rows = tuple(list_df.itertuples(index=False, name=None))
        # Creating a table object
        table = doc.add_table(rows=1, cols=len(list_df.columns))
        table.style = 'TableGrid'

        # Adding heading in the 1st row of the table
        row = table.rows[0].cells
        cols = list(list_df.columns)
        for i in range(len(cols)):
            row[i].text = cols[i]

        # Adding data from the list to the table
        for id, name, area in df_rows:
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(id)
            row[1].text = name
            row[2].text = area
        return doc
    except:
        print(f"Exception Occured:\n{traceback.format_exc()}")

def add_failed_tests(doc,df):
    try:
        # Generate list of tests failed
        test_failed_df = df[df["Result"] == 'Failed']
        test_failed_df = test_failed_df[["Test Id", "Test Name", "Failed Step", "Reasons", "Error Message"]]

        df_rows = tuple(test_failed_df.itertuples(index=False, name=None))
        # Creating a table object
        # doc.add_heading('List of tests failed', level=1)
        table = doc.add_table(rows=1, cols=len(test_failed_df.columns))
        table.style = 'TableGrid'

        # Adding heading in the 1st row of the table
        row = table.rows[0].cells

        cols = list(test_failed_df.columns)
        for i in range(len(cols)):
            row[i].text = cols[i]

        # Adding data from the list to the table

        for id, name, failed_step, reasons, error_message in df_rows:
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(id)
            row[1].text = str(name)
            row[2].text = str(failed_step)
            row[3].text = str(reasons)
            row[4].text = str(error_message)
        return doc
    except:
        print(f"Exception Occured:\n{traceback.format_exc()}")


try:
    '''
        Reads the input csv file as a dataframe and saves the data to word document as
        Table with list of Test Scenarios,
        Pie chart depicting Passed and Failed tests,
        Bar chart depicting Passed and Failed tests,
        Table with list of Failed tests.
    '''
    # Prompt the user for the CSV file path
    # csv_file_path = input("Enter the CSV file path: ")
    csv_file_path = 'C:/Users/LatheefS/Desktop/AT Python Files/x_ve1_maximise_test_results.csv'

    # Read the CSV file into a pandas DataFrame
    df = pd.read_csv(csv_file_path)
    df = df.fillna('')
    df.rename(columns={'number': 'Number',
                       'u_company': 'Company',
                       'error_message': 'Error Message',
                       'failed_step': 'Failed Step',
                       'process_id': 'Process Id',
                       'reasons': 'Reasons',
                       'result': 'Result',
                       'test_area': 'Test Area',
                       'test_duration': 'Test Duration',
                       'test_id': 'Test Id',
                       'test_name': 'Test Name'}, inplace=True)
    # Add tests list table to document
    doc = Document('C:/Users/LatheefS/Desktop/AT Python Files/AT_Report_Template.docx')

    Intro_heading = doc.add_heading('Introduction', level=1)
    Intro_heading.style.font.name = 'Verdana'
    Intro_heading.style.font.size = Pt(16)
    Intro_heading.style.font.bold = True
    Intro_heading.style.font.color.rgb = docx.shared.RGBColor(0, 182, 181)

    IntroPara = Static.static_dict['Introduction']
    In_Para = doc.add_paragraph(IntroPara)
    # In_Para.add_run("<Version>").bold = True
    # In_Para.add_run("<Version>").italic = True
    In_Para.paragraph_format.space_before = Pt(10)
    In_Para.paragraph_format.space_after = Pt(10)

    purpose_heading = doc.add_heading('Purpose ', level=2)
    purpose_heading.style.font.name = 'Verdana'
    purpose_heading.style.font.size = Pt(11)
    purpose_heading.style.font.bold = True
    purpose_heading.style.font.italic = True
    purpose_heading.style.font.color.rgb = docx.shared.RGBColor(9, 208, 162)

    PurposePara = Static.static_dict['Purpose']
    PPPara = doc.add_paragraph(PurposePara)
    PPPara.paragraph_format.space_before = Pt(10)
    PPPara.paragraph_format.space_after = Pt(10)

    PShortPara = Static.static_dict['Purpose_short']
    doc.add_paragraph(PShortPara)

    Scope_heading=doc.add_heading('Scope', level=1)
    # Scope_heading.style.font.name = 'Verdana'
    # Scope_heading.style.font.size = Pt(16)
    # Scope_heading.style.font.bold = True
    # Scope_heading.style.font.color.rgb = docx.shared.RGBColor(9, 208, 162)

    ScopePara = Static.static_dict['Scope']
    Sp_Para = doc.add_paragraph(ScopePara)
    # Sp_Para.paragraph_format.space_before = Pt(10)
    # Sp_Para.paragraph_format.space_after = Pt(10)

    doc.add_heading('Test Scenarios ', level=2)
    TestSPara = Static.static_dict['Tools']
    TSPara = doc.add_paragraph(TestSPara)
    TSPara.paragraph_format.space_before = Pt(10)
    TSPara.paragraph_format.space_after = Pt(10)

    doc.add_heading('Scenario Scope (Regression and Targeted Tests)', level=3)

    doc = add_test_lists(doc,df)
    pie_df = df.copy()
    pie_df = pie_df.apply(mark_others, axis=1)

    # Generate Pie Chart
    # Assuming your 'Result' column contains distinct values (e.g., 'Passed', 'Failed', etc.)
    result_counts = pie_df['Result'].value_counts()
    # Define custom colors for each value
    colors = {'Passed': '#7bdcb5','Failed': '#f78da7','Others': '#fcb900'}         # Add more colors for other values if needed
    # Create the donut pie chart
    plt.figure(figsize=(8, 6))
    # Calculate the total count and plot the center text (total count)
    total_count = result_counts.sum()
    plt.text(0, 0, f"{total_count}", fontsize=12, ha='center', va='center')
    # Create the pie chart
    wedges, _, _ = plt.pie(result_counts, labels=result_counts.index,
                           colors=[colors.get(val, '#abb8c3') for val in result_counts.index],
                           autopct='', pctdistance=0.85, wedgeprops={'edgecolor': 'w'})
    # Add individual counts on each wedge
    for wedge, count in zip(wedges, result_counts):
        angle = (wedge.theta2 - wedge.theta1) / 2.0 + wedge.theta1
        x = 0.85 * np.cos(np.deg2rad(angle))
        y = 0.85 * np.sin(np.deg2rad(angle))
        plt.text(x, y, str(count), ha='center', va='center')

    # Draw a white circle at the center to create the donut effect
    centre_circle = plt.Circle((0, 0), 0.50, fc='white')
    fig = plt.gcf()
    fig.gca().add_artist(centre_circle)

    # Add title and legend
    plt.title(r"$\bf{Release\ Overview}$", color='brown')
    plt.legend(result_counts.index, loc='upper right', prop={'weight': 'bold'})
    plt.axis('equal')  # Equal aspect ratio ensures a circular pie chart
    # plt.show() # To display during execution
    # Save the Pie chart to an image file
    piechart_image_path = 'pie_chart.png'
    plt.savefig(piechart_image_path)

    # Print the chart to a Word document
    chart_heading = doc.add_heading('Test Overview', level=1)
    # chart_heading.style.font.name = 'Verdana'
    # chart_heading.style.font.size = Pt(16)
    # chart_heading.style.font.bold = True
    # chart_heading.style.font.color.rgb = docx.shared.RGBColor(0, 182, 181)

    OverVPara = Static.static_dict['Test Overview']
    OV_Para = doc.add_paragraph(OverVPara)
    OV_Para.paragraph_format.space_before = Pt(10)
    OV_Para.paragraph_format.space_after = Pt(10)

    doc.add_picture(piechart_image_path, width=docx.shared.Inches(6))
    print("Pie Chart saved as 'pie_chart.png' and printed to 'AT Audit Report.docx'.")
    plt.close()

    # Generate Bar Chart
    color_map = {'Passed':'#7bdcb5','Failed':'#f78da7'}
    bar_df = df.copy()
    bar_df = bar_df[(bar_df['Result']=='Passed')|(bar_df['Result']=='Failed')]
    bargraph = sns.countplot(x='Test Area', hue='Result', data=bar_df, palette=color_map)
    sns.set(rc={'figure.figsize': (10, 10)})
    for i in bargraph.containers:
        bargraph.bar_label(i, )

    # Customize plot properties
    plt.title("Results By Test Area", color='brown', fontsize=14, weight='bold')
    plt.xlabel("Test Area")
    plt.ylabel("No.of Tests")
    plt.legend(loc='upper right', prop={'weight': 'bold'})
    # plt.show() # To display during execution

    # Save the Bar Graph to an image file
    bargraph_image_path = 'bar_chart.png'
    plt.savefig(bargraph_image_path)
    # Print the chart to a Word document
    doc.add_picture(bargraph_image_path, width=docx.shared.Inches(6))
    print("Bar Chart saved as 'bar_chart.png' and printed to 'AT Audit Report.docx'.")

    promptreply = PyAi.output1
    propmtprint=doc.add_paragraph(promptreply)
    propmtprint.paragraph_format.space_before = Pt(10)
    propmtprint.paragraph_format.space_after = Pt(10)

    # Add failed tests list table to document and save the document
    tf_heading=doc.add_heading('Tests not executed', level=2)
    tf_heading.style.font.name = 'Verdana'
    tf_heading.style.font.size = Pt(11)
    tf_heading.style.font.bold = True
    tf_heading.style.font.italic = True
    tf_heading.style.font.color.rgb = docx.shared.RGBColor(9, 208, 162)
    doc = add_failed_tests(doc, df)

    doc.add_heading('Tools and processes', level=2)
    ToolsPPara = Static.static_dict['Tools']
    ToolsPara = doc.add_paragraph(ToolsPPara)
    ToolsPara.paragraph_format.space_before = Pt(10)
    ToolsPara.paragraph_format.space_after = Pt(10)

    doc.add_heading('Environment set up and testing approach', level=2)
    EnvPara = Static.static_dict['Environment']
    EPara = doc.add_paragraph(EnvPara)
    EPara.paragraph_format.space_before = Pt(10)
    EPara.paragraph_format.space_after = Pt(10)

    doc.add_heading('Test execution entry criteria', level=2)
    EntryPara = Static.static_dict['Test Entry']
    EntPara = doc.add_paragraph(EntryPara)
    EntPara.paragraph_format.space_before = Pt(10)
    EntPara.paragraph_format.space_after = Pt(10)

    doc.add_heading('Test execution exit criteria', level=2)
    ExitPara = Static.static_dict['Test Exit']
    ExtPara = doc.add_paragraph(ExitPara)
    ExtPara.paragraph_format.space_before = Pt(10)
    ExtPara.paragraph_format.space_after = Pt(10)

    doc.add_heading('Evidence of Regression Test Completion', level=2)
    doc.add_heading('Functional testing', level=3)
    EviPara = Static.static_dict['Evidence']
    EvPara = doc.add_paragraph(EviPara)
    EvPara.paragraph_format.space_before = Pt(10)
    EvPara.paragraph_format.space_after = Pt(10)

    doc.add_heading('Defects', level=1)
    # chart_heading.style.font.name = 'Verdana'
    # chart_heading.style.font.size = Pt(16)
    # chart_heading.style.font.bold = True
    # chart_heading.style.font.color.rgb = docx.shared.RGBColor(0, 182, 181)
    DefectPara = Static.static_dict['Evidence']
    DfctPara = doc.add_paragraph(DefectPara)
    DfctPara.paragraph_format.space_before = Pt(10)
    DfctPara.paragraph_format.space_after = Pt(10)

    doc.add_heading('Failed tests', level=2)
    doc.add_heading('Defect Status Summary', level=2)

    doc.add_heading('Communications', level=1)
    CommPara = Static.static_dict['Communications']
    ComPara = doc.add_paragraph(CommPara)
    # ComPara.paragraph_format.space_before = Pt(10)
    # ComPara.paragraph_format.space_after = Pt(10)

    doc.add_heading('Lessons Learned', level=1)
    # LsnPara = Static.static_dict['']
    # LLPara = doc.add_paragraph(LsnPara)
    # # LLPara.paragraph_format.space_before = Pt(10)
    # # LLPara.paragraph_format.space_after = Pt(10)

    doc.add_heading('Issues and Recommended Actions', level=2)

    doc.save('AT Audit Report.docx')


except FileNotFoundError:
    print(f"File '{csv_file_path}' not found. Please check the path and try again.")
    exit()

except:
    print(f"Exception occurred:{traceback.format_exc()}")