import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

document = Document()
styles = document.styles
custom_para_style = styles.add_style('MyCustomParagraph', WD_STYLE_TYPE.PARAGRAPH)
custom_para_style.font.name = 'Calibri (Body)'  # Set the font name to 'Calibri'
custom_para_style.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # Red color
custom_para_style.font.size = Pt(16)
paragraph = document.add_paragraph(style=custom_para_style)
paragraph.add_run("This is a custom-styled paragraph.")


custom_heading_style = styles.add_style('MyCustomHeading', WD_STYLE_TYPE.PARAGRAPH)
custom_heading_style.font.size = Pt(16)  # Set font size to 16 points
custom_heading_style.font.bold = True  # Make the font bold

paragraph = document.add_paragraph(style=custom_heading_style)
paragraph.add_run("This is a custom-styled heading.")

style1=styles.add_style('MyCustomParagraph', WD_STYLE_TYPE.PARAGRAPH)
document.save('gfg.docx')