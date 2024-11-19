import docx
from docx import Document
from docx.shared import Pt, Inches
import Static

document = Document()
IntroHead = document.add_heading('Introduction')
IntroHead.style.font.name = 'Verdana'
IntroHead.style.font.size = Pt(16)
IntroHead.style.font.bold = True
IntroHead.style.font.color.rgb = docx.shared.RGBColor(0, 182, 181)

IntroPara = Static.static_dict['Test Overview']
paragraph = document.add_paragraph(IntroPara)
paragraph.paragraph_format.space_before = Pt(12)
paragraph.paragraph_format.space_after = Pt(10)
paragraph_e = document.add_paragraph('Lorem ipsum dolor sit amet.')

document.add_heading('Introduction',1)
paragraph_new = document.add_paragraph('Lorem ipsum dolor sit amets.')

document.add_heading('Introduction',2)
paragraph_new = document.add_paragraph('Lorem ipsum dolor sit amets.')
document.save('gfgh.docx')