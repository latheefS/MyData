import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import docx
import traceback

# Prompt the user for the CSV file path
#csv_file_path = input("Enter the CSV file path: ")
csv_file_path = 'C:/Users/LatheefS/Desktop/AT Python Files/x_ve1_maximise_test_results.csv'

# Ensure the file path ends with ".csv"
if not csv_file_path.lower().endswith(".csv"):
    csv_file_path += ".csv"


def mark_others(row):
    if row['result'] != 'Passed' and row['result'] != 'Failed' and row['result'] != 'N/A':
        row['result'] = 'Others'
    return row


# Read the CSV file into a pandas DataFrame
try:
    df = pd.read_csv(csv_file_path)
    df['result'] = df['result'].fillna('N/A')
    df = df.apply(mark_others, axis=1)
except FileNotFoundError:
    print(f"File '{csv_file_path}' not found. Please check the path and try again.")
    exit()

# Step 1: Read the CSV file
# csv_file_path = 'C:/Users/LatheefS/Desktop/AT Python Files/x_ve1_maximise_test_results.csv'
# df = pd.read_csv(csv_file_path)
# df.fillna('NA', inplace = True)
# df['result']=df['result'].fillna('N/A')
# def mark_others(row):
#     if row['result'] != 'Passed' and row['result'] != 'Failed' and row['result'] != 'N/A':
#         row['result'] = 'Others'
#     return row
#
# df = df.apply(mark_others,axis=1)
#
# print(df.to_string())
try:
    df = pd.read_csv(csv_file_path)
    df['result'] = df['result'].fillna('N/A')
    df = df.apply(mark_others, axis=1)

    # Step 2: Generate Pie Chart
    # Assuming your 'Result' column contains distinct values (e.g., 'Passed', 'Failed', etc.)
    result_counts = df['result'].value_counts()

    # Define custom colors for each value
    colors = {
        'Passed': '#7bdcb5',
        'Failed': '#f78da7',
        'Others': '#fcb900',
        # Add more colors for other values if needed
    }

    # Create the donut pie chart
    plt.figure(figsize=(8, 6))

    # Calculate the total count
    total_count = result_counts.sum()

    # Plot the center text (total count)
    plt.text(0, 0, f"{total_count}", fontsize=12, ha='center', va='center')

    # Create the pie chart
    wedges, _, _ = plt.pie(result_counts, labels=result_counts.index,
                           colors=[colors.get(val, '#abb8c3') for val in result_counts.index],
                           autopct='', pctdistance=0.85, wedgeprops={'edgecolor': 'w'})

    # Add individual counts on each wedge
    for wedge, count in zip(wedges, result_counts):
        angle = (wedge.theta2 - wedge.theta1) / 2.0 + wedge.theta1
        x = 0.85 * np.cos(np.deg2rad(angle))
        y = 0.85 * np.sin(np.deg2rad(angle))
        plt.text(x, y, str(count), ha='center', va='center')

    # Draw a white circle at the center to create the donut effect
    centre_circle = plt.Circle((0, 0), 0.50, fc='white')
    fig = plt.gcf()
    fig.gca().add_artist(centre_circle)

    # Add title and legend
    plt.title(r"$\bf{Release\ Overview}$", color='brown')
    plt.legend(result_counts.index, loc='upper right', prop={'weight': 'bold'})

    # Display the plot
    plt.axis('equal')  # Equal aspect ratio ensures a circular pie chart
    # plt.show()

    # Save the Pie chart to an image file
    piechart_image_path = 'pie_chart.png'
    plt.savefig(piechart_image_path)

    # Step 4: Print the chart to a Word document
    doc = Document()
    doc.add_heading('3.Test Overview', level=1)
    doc.add_picture(piechart_image_path, width=docx.shared.Inches(6))
    # doc.save('AT Completion Report.docx')

    print("Pie Chart saved as 'pie_chart.png' and printed to 'AT Completion Report.docx'.")
    plt.close()


    # Step 3: Generate Bar Chart
    unique_results = list(set(list(df['result'])))
    color_map = {}
    for val in unique_results:
        if val == 'Passed':
            color_map[val] = "#7bdcb5"
        elif val == 'Failed':
            color_map[val] = "#f78da7"
        elif val == 'Others':
            color_map[val] = "#fcb900"
        else:
            color_map[val] = "#abb8c3"

    bargraph = sns.countplot(x='test_area', hue='result', data=df, palette=color_map)
    sns.set(rc={'figure.figsize': (10, 10)})
    for i in bargraph.containers:
        bargraph.bar_label(i, )

    # Customize plot properties
    plt.title("Results By Test Area", color='brown', fontsize=14, weight='bold')
    plt.xlabel("Test Area")
    plt.ylabel("No.of Tests")
    plt.legend(loc='upper right', prop={'weight': 'bold'})

    # Show the plot
    # plt.show()

    # Save the Bar Graph to an image file
    bargraph_image_path = 'bar_chart.png'
    plt.savefig(bargraph_image_path)

    # Step 4: Print the chart to a Word document
    #doc = Document()
    #doc.add_heading('Bar Graph from CSV Data', level=1)
    doc.add_picture(bargraph_image_path, width=docx.shared.Inches(6))
    #doc.save('AT Completion Report.docx')

    print("Bar Chart saved as 'bar_chart.png' and printed to 'AT Completion Report.docx'.")

    '''
    # Generate list of tests
    test_list_df = df[["test_id", "test_name", "test_area"]]
    print(test_list_df.to_string())

    # Generate list of tests passed
    test_passed_df = df[df["result"] == 'Passed']
    test_passed_df = test_passed_df[["test_id", "test_name", "test_area", "test_duration", "result"]]
    print(test_passed_df.to_string())

    # Generate list of tests failed
    test_failed_df = df[df["result"] == 'Failed']
    test_failed_df = test_failed_df[
        ["test_id", "test_name", "test_area", "failed_step", "reasons", "error_message", "result"]]
    print(test_failed_df.to_string())
    '''

except FileNotFoundError:
    print(f"File '{csv_file_path}' not found. Please check the path and try again.")
    exit()

except:
    print(f"Exception occurred:{traceback.format_exc()}")

# except Exception as e:
#     print(f"Exception occurred:\n{e}")
#     exit()
##########---------------------------------------------###################
# Add a Title to the document
doc.add_heading('List of tests passed', level=1)

# Table data in a form of list
df = pd.read_csv(csv_file_path)
list_df = df[["test_id", "test_name", "test_area"]]
df_rows=tuple(list_df.itertuples(index=False, name=None))
# Creating a table object
table = doc.add_table(rows=1, cols=len(list_df.columns))
table.style = 'TableGrid'

# Adding heading in the 1st row of the table
row = table.rows[0].cells
cols=list(list_df.columns)
for i in range(len(cols)):
   row[i].text = cols[i]


# Adding data from the list to the table

for id, name, area in df_rows:

   # Adding a row and then adding data in it.
   row = table.add_row().cells
   # Converting id to string as table can only take string input
   row[0].text = str(id)
   row[1].text = name
   row[2].text = area



# Generate list of tests passed

test_passed_df = df[df["result"] == 'Passed']
test_passed_df = test_passed_df[["test_id", "test_name", "test_area", "test_duration", "result"]]


df_rows=tuple(test_passed_df.itertuples(index=False, name=None))
# Creating a table object

doc.add_heading('List of tests passed', level=1)
table = doc.add_table(rows=1, cols=len(test_passed_df.columns))
table.style = 'TableGrid'

# Adding heading in the 1st row of the table
row = table.rows[0].cells


cols=list(test_passed_df.columns)
for i in range(len(cols)):
   row[i].text = cols[i]


# Adding data from the list to the table

for id, name, area, duration, result in df_rows:

   # Adding a row and then adding data in it.
   row = table.add_row().cells
   # Converting id to string as table can only take string input
   row[0].text = str(id)
   row[1].text = str(name)
   row[2].text = str(area)
   row[3].text = str(duration)
   row[4].text = str(result)

# Generate list of tests failed
test_failed_df = df[df["result"] == 'Failed']
test_failed_df = test_failed_df[["test_id", "test_name", "test_area", "failed_step", "reasons", "error_message", "result"]]

df_rows=tuple(test_failed_df.itertuples(index=False, name=None))
# Creating a table object
doc.add_heading('List of tests failed', level=1)
table = doc.add_table(rows=1, cols=len(test_failed_df.columns))
table.style = 'TableGrid'


# Adding heading in the 1st row of the table
row = table.rows[0].cells

cols=list(test_failed_df.columns)
for i in range(len(cols)):
   row[i].text = cols[i]


# Adding data from the list to the table

for id, name, area, failed_step, reasons, error_message, result in df_rows:

   # Adding a row and then adding data in it.
   row = table.add_row().cells
   # Converting id to string as table can only take string input
   row[0].text = str(id)
   row[1].text = str(name)
   row[2].text = str(area)
   row[3].text = str(failed_step)
   row[4].text = str(reasons)
   row[5].text = str(error_message)
   row[6].text = str(result)
# Now save the document to a location
doc.save('gfg.docx')