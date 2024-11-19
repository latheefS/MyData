import docx
from docx import Document
from docx.enum.style import WD_STYLE
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

document = Document()
# styles = document.styles
# # style = styles[WD_STYLE.BODY_TEXT]
#
#
#
# # Create a new heading style based on 'Heading 1'
# new_heading_style = document.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
# new_heading_style.base_style = document.styles['Heading 1']
#
# # Modify the font properties of the new style
# new_heading_style.font.name = 'Verdana'
# new_heading_style.font.size = Pt(16)  # Set the desired font size
# new_heading_style.font.bold = True
#
# # document.add_heading('Introduction')
# document.add_heading('Introduction',1)

heading = document.add_heading('1.My Custom Heading', level=1)


document.save('ffg.docx')