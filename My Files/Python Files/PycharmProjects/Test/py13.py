from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Define the sample paragraphs
paragraph1 = "This is a sample paragraph. Some words will be bold."
paragraph2 = "This is another sample paragraph. Some words will be italic."

# Create a new Word document
doc = Document()

# Add the paragraphs to the document
doc.add_paragraph(paragraph1)
doc.add_paragraph(paragraph2)

# Function to format specific words in a paragraph
def format_words(paragraph, words, style):
    p = paragraph._element
    for word in words:
        r = p.get_or_add_r()
        t = r.get_or_add_t()
        t.text = word
        rPr = r.get_or_add_rPr()
        rPr.append(parse_xml(r'<w:' + style + ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="1"/>', nsdecls('w')))

# Words to be formatted in paragraph1
bold_words = ["sample", "bold"]
format_words(doc.paragraphs[0], bold_words, "b")

# Words to be formatted in paragraph2
italic_words = ["another", "italic"]
format_words(doc.paragraphs[1], italic_words, "i")

# Save the document
doc.save("formatted_paragraphs.docx")

print("Formatted paragraphs saved to formatted_paragraphs.docx")

