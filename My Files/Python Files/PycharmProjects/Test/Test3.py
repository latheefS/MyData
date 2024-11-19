import docx
from docx import Document

# Create a document
doc = Document()

# Add a paragraph with your text
paraText = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Vestibulum volutpat nulla sit amet lacinia sapien."

# Identify the word(s) to make bold
string2Bold = "Vestibulum volutpat"

# Add the paragraph to the document
paragraph = doc.add_paragraph(paraText, style=None)

# Split the paragraph into substrings
for oldPara in doc.paragraphs:
    if oldPara.text.find(string2Bold) >= 0:
        # Insert an empty paragraph before the existing one
        newPara = oldPara.insert_paragraph_before(text=None, style=None)
        # Add the substrings as separate runs
        for word in oldPara.text.split():
            run = newPara.add_run(word)
            if word == string2Bold:
                run.bold = True

# Save the document
doc.save("output.docx")
