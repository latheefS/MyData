import docx

# Create a new Word document
doc = docx.Document()

# Your paragraph with placeholders
IntroPara = "This document details the results of the <Version> Quarterly Update Testing. As per Oracle’s timetable for quarterly releases for wave <number>, the upgrade took place on non-production environments on the <day(dd/mm/yy)>, and into the production environment on the <day(dd/mm/yy)>."

# Add the paragraph to the document
doc.add_paragraph(IntroPara)

# Iterate through the runs in the paragraph
for run in doc.paragraphs[0].runs:
    if "<Version>" in run.text:
        run.bold = True
    elif "<number>" in run.text:
        run.bold = True
    elif "<day(dd/mm/yy)>" in run.text:
        run.bold = True

# Save the document
doc.save("formatted_paragraphs.docx")

print("Formatted paragraph saved as formatted_paragraphs.docx")

