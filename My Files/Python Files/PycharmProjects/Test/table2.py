from docx import Document
from docx.shared import RGBColor, Pt

# Create a new Word document
doc = Document()

# Add Table 1 (3x3)
table1 = doc.add_table(rows=3, cols=3, style="Table Grid")
column_headers1 = ["Criteria", "Met / not met", "Mitigation factors (if not met)"]
header_row1 = table1.rows[0]
for col_idx, header_text in enumerate(column_headers1):
    cell = header_row1.cells[col_idx]
    cell.text = header_text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black color
    # shading_elm = cell._tc.get_or_add_tcPr().add_shd()
    # shading_elm.val = "clear"
    # shading_elm.color = "auto"
    # shading_elm.fill = "808080"  # Grey color

# # Add a page break
# doc.add_page_break()

# Add Table 2 (4x4)
table2 = doc.add_table(rows=4, cols=4, style="Table Grid")
column_headers2 = ["Env Name", "Config Requirements", "Start", "End"]
header_row2 = table2.rows[0]
for col_idx, header_text in enumerate(column_headers2):
    cell = header_row2.cells[col_idx]
    cell.text = header_text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black color
    # shading_elm = cell._tc.get_or_add_tcPr().add_shd()
    # shading_elm.val = "clear"
    # shading_elm.color = "auto"
    # shading_elm.fill = "808080"  # Grey color

# Save the document
doc.save("custom_tables.docx")

print("Custom tables with specified header formatting saved as custom_tables.docx")
