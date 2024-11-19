from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Define the paragraphs
paragraph1 = "This is a sample paragraph. Some words will be bold."
paragraph2 = "This is another sample paragraph. Some words will be italic."

# Create a new Word document
doc = Document()

# Add the paragraphs to the document
doc.add_paragraph(paragraph1)
doc.add_paragraph(paragraph2)

# Access the paragraphs
p1 = doc.paragraphs[0]
p2 = doc.paragraphs[1]

# Add bold formatting to specific words in paragraph1
for run in p1.runs:
    if "bold" in run.text.lower():
        run.bold = True

# Add italic formatting to specific words in paragraph2
for run in p2.runs:
    if "italic" in run.text.lower():
        run.italic = True

# Save the document
doc.save("formatted_paragraphs.docx")

# Print a success message
print("Formatted paragraphs saved to formatted_paragraphs.docx")

