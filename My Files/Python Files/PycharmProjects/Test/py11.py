from docx import Document
from docx.shared import RGBColor, Pt

# Create a new Word document
doc = Document()

# Add a table with 3 rows and 3 columns
table = doc.add_table(rows=3, cols=3, style="Table Grid")

# Set the column headers
column_headers = ["Criteria", "Met / not met", "Mitigation factors (if not met)"]

# Set the header row background color to grey
header_row = table.rows[0]
for col_idx, header_text in enumerate(column_headers):
    cell = header_row.cells[col_idx]
    cell.text = header_text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black color
    # shading_elm = cell._tc.get_or_add_tcPr().add_shd()
    # shading_elm.val = "clear"
    # shading_elm.color = "auto"
    # shading_elm.fill = "808080"  # Grey color

# Save the document
doc.save("empty_table.docx")

print("Custom table with specified header formatting saved as empty_table.docx")
